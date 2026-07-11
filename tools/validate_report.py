from pathlib import Path
from zipfile import ZipFile
from docx import Document

p = Path("outputs/软件综合实践报告-人工智能1班-20240282-杨杨文琦-完成版.docx")
d = Document(p)
with ZipFile(p) as z:
    error = z.testzip()
    images = sum(name.startswith("word/media/") for name in z.namelist())
print(
    f"paragraphs={len(d.paragraphs)} tables={len(d.tables)} "
    f"sections={len(d.sections)} images={images} zip_error={error} size={p.stat().st_size}"
)
