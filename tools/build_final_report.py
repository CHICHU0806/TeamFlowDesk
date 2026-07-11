from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "软件综合实践报告-人工智能1班-20240282-杨杨文琦-完成版.docx"
ASSET = ROOT / "Assets" / "TeamFlowDeskIcon.png"
FIG = OUT / "report_figures"
FIG.mkdir(exist_ok=True)

BLUE = "1F4E79"
LIGHT = "D9EAF7"
GRAY = "F2F4F7"


def font(size=10.5, bold=False, color="000000", name="宋体"):
    return {"size": Pt(size), "bold": bold, "color": RGBColor.from_string(color), "name": name}


def apply_run(run, **kw):
    cfg = font(**kw)
    run.font.size = cfg["size"]
    run.font.bold = cfg["bold"]
    run.font.color.rgb = cfg["color"]
    run.font.name = cfg["name"]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cfg["name"])


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    apply_run(r, size={1: 16, 2: 14, 3: 12}[level], bold=True, color=BLUE, name="黑体")
    return p


def add_body(doc, text, first=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if first else None
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    apply_run(r)
    return p


def add_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(3)
        apply_run(p.add_run(item))


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run(p.add_run(h), size=9.5, bold=True, color="FFFFFF", name="黑体")
        if widths:
            cell.width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2:
                set_cell_shading(cells[i], GRAY)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            apply_run(p.add_run(str(value)), size=9)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def text_font(size, bold=False):
    candidates = [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf"]
    for f in candidates:
        if Path(f).exists():
            return ImageFont.truetype(f, size=size, index=0)
    return ImageFont.load_default()


def make_architecture(path):
    im = Image.new("RGB", (1500, 820), "white")
    d = ImageDraw.Draw(im)
    title = text_font(42, True)
    body = text_font(30)
    d.text((540, 35), "TeamFlowDesk 系统架构", fill="#163A5F", font=title)
    layers = [
        ("表示层", "WinUI 3 页面、导航、对话框、状态提示", "#D9EAF7"),
        ("业务层", "项目、任务、成员、设备、周报、AI 协作逻辑", "#E5F3EA"),
        ("数据访问层", "仓储服务、参数化 SQL、事务与备份", "#FFF1D6"),
        ("数据层", "SQLite 本地数据库", "#F3E6F7"),
    ]
    y = 120
    for name, desc, color in layers:
        d.rounded_rectangle((190, y, 1310, y + 130), radius=18, fill=color, outline="#56738F", width=3)
        d.text((245, y + 37), name, fill="#163A5F", font=title)
        d.text((520, y + 47), desc, fill="#263746", font=body)
        y += 165
    im.save(path)


def make_workflow(path):
    im = Image.new("RGB", (1600, 700), "white")
    d = ImageDraw.Draw(im)
    title = text_font(42, True)
    body = text_font(28)
    d.text((590, 35), "任务闭环工作流", fill="#163A5F", font=title)
    nodes = [("待处理", "录入任务\n设置负责人和截止时间"), ("进行中", "开始执行\n持续更新任务状态"), ("待验收", "提交成果\n等待项目负责人确认"), ("已完成", "验收通过\n进入统计与周报")]
    xs = [70, 455, 840, 1225]
    for i, ((name, sub), x) in enumerate(zip(nodes, xs)):
        d.rounded_rectangle((x, 175, x + 300, 475), radius=22, fill="#EAF2F8", outline="#1F4E79", width=4)
        d.text((x + 85, 215), name, fill="#1F4E79", font=title)
        for j, line in enumerate(sub.split("\n")):
            d.text((x + 35, 315 + 46*j), line, fill="#334455", font=body)
        if i < 3:
            d.line((x + 305, 325, xs[i+1] - 18, 325), fill="#1F4E79", width=7)
            d.polygon([(xs[i+1]-18, 325), (xs[i+1]-45, 310), (xs[i+1]-45, 340)], fill="#1F4E79")
    d.text((540, 545), "异常分支：任务可退回上一状态，并保留修正后的结果", fill="#A33B20", font=body)
    im.save(path)


def make_modules(path):
    im = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(im)
    title = text_font(40, True)
    body = text_font(27)
    d.rounded_rectangle((520, 45, 980, 145), radius=18, fill="#1F4E79")
    d.text((625, 72), "TeamFlowDesk", fill="white", font=title)
    mods = ["工作台", "项目管理", "任务管理", "成员管理", "设备管理", "AI 协作记录", "周报中心", "系统设置"]
    for i, name in enumerate(mods):
        col, row = i % 4, i // 4
        x, y = 70 + col * 360, 260 + row * 260
        d.line((750, 145, x + 145, y), fill="#7E9BB5", width=3)
        d.rounded_rectangle((x, y, x + 290, y + 145), radius=16, fill="#EAF2F8", outline="#1F4E79", width=3)
        tw = d.textbbox((0,0), name, font=body)[2]
        d.text((x + (290-tw)/2, y + 53), name, fill="#183B5B", font=body)
    im.save(path)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_run(p.add_run(text), size=9, color="555555")


def build():
    arch = FIG / "architecture.png"
    flow = FIG / "workflow.png"
    modules = FIG / "modules.png"
    make_architecture(arch); make_workflow(flow); make_modules(modules)

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21)
    sec.top_margin, sec.bottom_margin = Cm(2.5), Cm(2.2)
    sec.left_margin, sec.right_margin = Cm(2.7), Cm(2.4)
    add_page_number(sec.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"; normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for level in (1, 2, 3):
        s = styles[f"Heading {level}"]
        s.font.name = "黑体"; s.font.color.rgb = RGBColor.from_string(BLUE)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        s.paragraph_format.space_before = Pt(10); s.paragraph_format.space_after = Pt(6)

    # Cover
    if ASSET.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSET), width=Cm(2.6))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    apply_run(p.add_run("软件综合实践报告"), size=28, bold=True, color=BLUE, name="黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_run(p.add_run("TeamFlowDesk 团队运行管理工作流系统"), size=18, bold=True, name="黑体")
    doc.add_paragraph("\n\n")
    cover = [("院　　系", "人工智能学院"), ("专　　业", "人工智能"), ("班　　级", "人工智能1班"), ("姓　　名", "杨杨文琦"), ("学　　号", "20240282"), ("项目名称", "TeamFlowDesk 团队运行管理工作流系统"), ("完成时间", "2026年7月")]
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in cover:
        cells = table.add_row().cells
        cells[0].width, cells[1].width = Cm(4), Cm(8)
        for c in cells: set_cell_margins(c, 140, 160, 140, 160)
        p0 = cells[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        apply_run(p0.add_run(k + "："), size=12, bold=True, name="黑体")
        apply_run(cells[1].paragraphs[0].add_run(v), size=12)
    doc.add_page_break()

    add_heading(doc, "摘　要", 1)
    add_body(doc, "TeamFlowDesk 是一套面向小型学生团队和实训项目组的桌面工作流管理系统。系统以项目为组织边界，以任务状态流转为核心，将成员、设备、周报和 AI 协作记录纳入统一的数据体系，解决信息分散、责任不清、进度难追踪和成果难沉淀等问题。项目采用 C#、.NET 10、WinUI 3、Windows App SDK 与 SQLite 实现，支持项目及任务全生命周期管理、成员负载查看、设备借用登记、周报生成与导出、数据库备份、紧凑导航和本地发布运行。开发过程中通过多轮人机协作完成需求澄清、方案比较、故障定位和界面改进，并对 AI 建议进行了代码验证与实际运行验证。最终系统形成了从任务建立、执行、验收到统计归档的完整闭环，可通过独立 EXE 在 Windows 环境运行。")
    add_body(doc, "关键词：团队工作流；任务管理；WinUI 3；SQLite；AI 协作；桌面应用", first=False)
    doc.add_page_break()

    add_heading(doc, "目　录", 1)
    toc = ["1 绪论", "2 需求分析", "3 数据库设计", "4 系统设计", "5 编码实现", "6 程序调试情况", "7 系统总结", "8 课程设计总结", "9 参考文献"]
    for i, t in enumerate(toc, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
        apply_run(p.add_run(t), size=12, bold=True if i < 9 else False, name="黑体" if i < 9 else "宋体")
    doc.add_page_break()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_body(doc, "在软件综合实践中，团队通常同时处理项目计划、任务分配、成员协作、设备使用和阶段汇报。若依赖聊天记录与多个零散表格，信息容易重复或遗漏，任务负责人和完成标准也不够清晰。尤其在短周期实训中，项目管理工具必须轻量、直观且能够快速落地，因此本项目选择开发一套本地桌面工作流系统。")
    add_heading(doc, "1.2 项目目标", 2)
    add_body(doc, "本项目目标是建立一个可独立运行、数据可持续保存、操作流程完整的团队运行管理系统。系统不仅记录任务，还通过状态流转、截止时间提醒、项目关联和周报汇总推动任务真正完成；同时保留 AI 协作过程，形成可复盘的开发记录。")
    add_list(doc, ["建立项目、任务、成员、设备、AI 记录和周报的统一数据入口。", "实现待处理、进行中、待验收、已完成的任务闭环。", "提供搜索、筛选、统计、导出和备份能力。", "采用符合 Windows 桌面应用习惯的界面与导航方式。", "发布自包含 EXE，使系统脱离开发工具也能运行。"])
    add_heading(doc, "1.3 技术路线与工作方法", 2)
    add_body(doc, "项目采用分层思路组织代码：WinUI 3 负责界面展示与交互，业务逻辑负责状态校验和统计，数据访问层通过参数化 SQL 读写 SQLite。开发流程按照“需求分析—原型实现—功能补齐—异常修复—界面优化—发布验证”推进。AI 用于辅助分析和生成候选方案，最终结论均通过阅读官方接口、编译结果和程序运行行为进行验证。")

    add_heading(doc, "2 需求分析", 1)
    add_heading(doc, "2.1 用户与应用场景", 2)
    add_body(doc, "系统主要面向 3 至 10 人的小型项目团队。项目负责人关注整体进度、任务逾期和成员分工；普通成员关注自己的待办、任务状态和设备使用情况；团队全体需要通过周报了解阶段成果。由于使用场景集中在个人电脑和小型团队，不需要复杂服务器部署，本地数据库能够降低使用门槛。")
    add_heading(doc, "2.2 功能需求", 2)
    add_table(doc, ["编号", "模块", "主要功能", "完成情况"], [
        ("F01", "工作台", "展示项目、任务、成员、逾期数量及近期动态", "已完成"),
        ("F02", "项目管理", "项目新增、编辑、删除、状态和进度维护", "已完成"),
        ("F03", "任务管理", "新增、搜索、筛选、关联项目、负责人、截止日期和状态流转", "已完成"),
        ("F04", "成员管理", "成员档案、角色信息及工作负载查看", "已完成"),
        ("F05", "设备管理", "设备登记、状态维护及借用信息管理", "已完成"),
        ("F06", "AI 协作", "记录问题、AI 建议、人工判断和修正结果", "已完成"),
        ("F07", "周报中心", "汇总阶段数据并导出 TXT 报告", "已完成"),
        ("F08", "系统设置", "数据库备份、关于信息和运行数据说明", "已完成"),
    ], [1.2, 2.5, 9.5, 2.2])
    add_heading(doc, "2.3 非功能需求", 2)
    add_list(doc, ["易用性：主导航清晰，侧栏收起后仍通过图标识别功能。", "可靠性：输入数据进行校验，数据库初始化和访问异常可提示。", "性能：本地常规数据规模下页面切换和查询响应及时。", "可维护性：界面、业务与数据访问职责分离，模块命名统一。", "可交付性：生成自包含 Windows 可执行版本，提供数据库备份。"])
    add_heading(doc, "2.4 系统功能模块", 2)
    doc.add_picture(str(modules), width=Cm(15.8)); caption(doc, "图2-1 系统功能模块图")
    add_heading(doc, "2.5 核心用例说明", 2)
    add_table(doc, ["用例", "参与者", "前置条件", "基本流程", "结果"], [
        ("创建项目", "项目负责人", "进入项目页", "填写名称、说明和日期后保存", "生成项目记录"),
        ("创建任务", "团队成员", "存在可选项目和成员", "填写标题、项目、负责人、优先级和截止时间", "任务进入待处理"),
        ("推进任务", "负责人", "任务已存在", "依次推进到进行中、待验收和已完成", "统计数据同步变化"),
        ("生成周报", "项目负责人", "已有本周项目数据", "系统汇总任务和成员信息并导出", "形成可归档文本"),
        ("备份数据", "系统用户", "数据库已初始化", "选择备份功能并确认保存", "生成数据库副本"),
    ], [2.2, 2, 3.3, 6.2, 2.5])

    add_heading(doc, "3 数据库设计", 1)
    add_heading(doc, "3.1 数据流分析", 2)
    add_body(doc, "用户在各功能页面输入业务数据，界面层完成基础格式校验后交由服务层处理。服务层根据业务规则执行新增、修改、删除、状态推进或统计，再由数据访问层以参数化 SQL 写入 SQLite。工作台和周报中心读取多个业务表并形成聚合结果；备份功能复制完整数据库文件，从而实现数据恢复基础。")
    add_heading(doc, "3.2 概念结构设计", 2)
    add_body(doc, "核心实体包括项目、任务、成员、设备、AI 协作记录和周报。一个项目可包含多个任务，一个成员可负责多个任务；设备借用信息与成员关联；周报对特定时间段内的项目和任务数据进行汇总；AI 协作记录独立保存问题、建议、人工验证和最终结论。")
    add_heading(doc, "3.3 数据字典", 2)
    add_table(doc, ["数据表", "关键字段", "字段说明"], [
        ("Projects", "Id, Name, Description, Status, Progress, StartDate, EndDate", "项目基本信息、状态与总体进度"),
        ("Tasks", "Id, ProjectId, Title, AssigneeId, Priority, Status, DueDate", "任务归属、负责人、优先级和截止日期"),
        ("Members", "Id, Name, Role, Email, Phone, Status", "成员身份、联系方式及在岗状态"),
        ("Equipment", "Id, Name, Code, Status, BorrowerId, BorrowDate", "设备编码、可用状态与借用信息"),
        ("AiRecords", "Id, Question, AiSuggestion, HumanDecision, Correction, CreatedAt", "AI 建议、人工判断及纠正记录"),
        ("WeeklyReports", "Id, WeekStart, WeekEnd, Summary, CreatedAt", "周报时间范围与汇总内容"),
    ], [2.6, 8, 5.7])
    add_heading(doc, "3.4 逻辑结构与约束", 2)
    add_body(doc, "各数据表均使用整数主键，日期统一保存为可排序的文本格式。Tasks.ProjectId 关联 Projects.Id，Tasks.AssigneeId 关联 Members.Id。删除项目时系统先检查关联任务，避免出现失去归属的业务数据。任务状态限定在待处理、进行中、待验收和已完成范围内，优先级限定为低、中、高。")
    add_heading(doc, "3.5 主要建表语句", 2)
    sql = """CREATE TABLE IF NOT EXISTS Projects (
    Id INTEGER PRIMARY KEY AUTOINCREMENT,
    Name TEXT NOT NULL, Description TEXT,
    Status TEXT NOT NULL, Progress INTEGER NOT NULL DEFAULT 0,
    StartDate TEXT, EndDate TEXT, CreatedAt TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS Tasks (
    Id INTEGER PRIMARY KEY AUTOINCREMENT,
    ProjectId INTEGER, Title TEXT NOT NULL, Description TEXT,
    AssigneeId INTEGER, Priority TEXT NOT NULL, Status TEXT NOT NULL,
    DueDate TEXT, CreatedAt TEXT NOT NULL, UpdatedAt TEXT NOT NULL,
    FOREIGN KEY(ProjectId) REFERENCES Projects(Id),
    FOREIGN KEY(AssigneeId) REFERENCES Members(Id)
);"""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.right_indent = Cm(0.7)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    setp = OxmlElement("w:shd"); setp.set(qn("w:fill"), "F3F5F7"); p._p.get_or_add_pPr().append(setp)
    apply_run(p.add_run(sql), size=8.5, name="Consolas")
    add_body(doc, "实际初始化还包含 Members、Equipment、AiRecords 和 WeeklyReports 表。所有写入语句使用参数传值，避免字符串拼接引起格式错误或注入风险；数据库文件随应用数据保存，并可由设置页一键备份。")

    add_heading(doc, "4 系统设计", 1)
    add_heading(doc, "4.1 总体架构", 2)
    doc.add_picture(str(arch), width=Cm(15.8)); caption(doc, "图4-1 TeamFlowDesk 分层架构图")
    add_body(doc, "表示层由 MainWindow 和各业务 Page 组成，负责导航、表单、列表和反馈；业务层负责任务状态规则、统计和导出；数据访问层统一管理 SQLite 连接、查询和事务；数据层保存所有业务实体。该结构使界面调整不会直接影响数据库结构，也便于定位异常。")
    add_heading(doc, "4.2 页面与导航设计", 2)
    add_body(doc, "主窗口采用左侧 NavigationView。展开时显示图标和完整文字，收起时仅保留对应图标，避免原先只显示半截文字的问题。左上角使用千里马头作为专用应用图标，启动阶段先执行淡入、短暂停留和淡出动画，再进入主页面。工作台作为默认入口，业务模块按项目、任务、资源、记录和设置顺序排列。")
    add_heading(doc, "4.3 任务状态设计", 2)
    doc.add_picture(str(flow), width=Cm(16)); caption(doc, "图4-2 任务闭环状态流转图")
    add_body(doc, "任务状态不是普通标签，而是业务流程控制。新建任务默认为待处理，开始执行后进入进行中，提交成果后进入待验收，负责人确认后进入已完成。对于误操作或验收不通过情况，任务可回退并重新推进。逾期判定由截止日期、当前时间和完成状态共同决定。")
    add_heading(doc, "4.4 关键类与职责", 2)
    add_table(doc, ["类别", "代表对象", "主要职责"], [
        ("窗口与导航", "MainWindow", "启动动画、标题栏、侧栏导航和页面切换"),
        ("业务页面", "DashboardPage、TasksPage 等", "展示数据、采集输入、触发操作并提示结果"),
        ("实体模型", "Project、TaskItem、Member 等", "描述业务数据和界面绑定属性"),
        ("数据服务", "DatabaseService", "初始化数据库、执行查询和持久化操作"),
        ("辅助服务", "Report/Backup 相关逻辑", "生成报告、导出文本和复制数据库文件"),
    ], [2.8, 5, 8.5])
    add_heading(doc, "4.5 典型时序", 2)
    add_body(doc, "以创建任务为例：用户在 TasksPage 打开新建对话框并填写信息；页面校验标题、项目、负责人和日期；校验通过后调用数据服务执行参数化 INSERT；服务返回结果后页面重新加载任务列表和统计信息；界面显示成功提示。若数据库操作失败，则捕获异常并提示原因，不更新当前列表。")

    add_heading(doc, "5 编码实现", 1)
    add_heading(doc, "5.1 开发与运行环境", 2)
    add_table(doc, ["项目", "配置"], [("开发语言", "C#"), ("运行框架", ".NET 10"), ("桌面框架", "WinUI 3 / Windows App SDK"), ("数据库", "SQLite"), ("开发工具", "JetBrains Rider"), ("运行平台", "Windows 10/11 x64"), ("发布方式", "Self-contained 独立 EXE")], [4.2, 12.2])
    add_heading(doc, "5.2 主窗口与开屏动画", 2)
    add_body(doc, "程序启动后显示居中的千里马头图标，通过透明度动画逐渐显现，短暂停留后逐渐消失，再显示主窗口内容。动画总时长约 1.25 秒，不阻塞数据库初始化。主窗口同步设置应用标题和图标，使任务栏、窗口标题栏与系统品牌一致。")
    if ASSET.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSET), width=Cm(4.2)); caption(doc, "图5-1 TeamFlowDesk 千里马专用图标")
    add_heading(doc, "5.3 工作台实现", 2)
    add_body(doc, "工作台通过聚合查询获取项目数、进行中任务数、逾期任务数和成员数，并展示近期任务。统计卡片保持紧凑、可扫描，颜色只用于状态区分。工作台不承担复杂编辑，用户可通过导航进入对应模块处理详细业务。")
    add_heading(doc, "5.4 项目与任务管理实现", 2)
    add_body(doc, "项目页支持新增、编辑、删除及进度维护。任务页是系统核心，提供关键词搜索、项目筛选、状态筛选、任务详情、截止日期预警和状态推进。页面初始化时通过就绪标记避免 SelectionChanged 在控件尚未完成加载时触发查询，从而解决进入任务页崩溃的问题。刷新列表时保留筛选条件，减少重复操作。")
    add_heading(doc, "5.5 成员与设备管理实现", 2)
    add_body(doc, "成员模块保存姓名、角色、联系方式和状态，并结合任务数据计算成员负载。设备模块使用唯一编码区分设备，记录可用、借出等状态以及借用人和日期。两类资源均采用对话框完成新增和编辑，删除前进行确认，避免误操作。")
    add_heading(doc, "5.6 AI 协作记录实现", 2)
    add_body(doc, "AI 协作模块用于记录开发中提出的问题、AI 给出的候选建议、人工判断、验证方式和最终修正。它不是简单聊天记录，而是将“建议—验证—纠错—结论”形成可检索证据。例如对任务页崩溃、XAML 资源路径和 SQLite 依赖告警等问题，均保留了人工复核后的处理结果。")
    add_heading(doc, "5.7 周报、导出与备份", 2)
    add_body(doc, "周报中心按时间范围汇总项目和任务信息，可生成便于提交和归档的 TXT 文本。设置页提供数据库备份，系统在用户文档目录下创建 TeamFlowDeskBackups 文件夹，并按时间命名数据库副本。导出和备份均返回明确的成功或失败信息。")
    add_heading(doc, "5.8 发布运行", 2)
    add_body(doc, "项目采用自包含方式发布为 Windows x64 可执行程序，运行目录包含应用所需组件和资源。用户可直接双击 TeamFlowDesk.exe 启动，不再依赖 Rider，也不要求目标电脑预先安装对应 .NET 运行时。发布后对启动动画、导航、任务页和数据库读写进行了复测。")

    add_heading(doc, "6 程序调试情况", 1)
    add_heading(doc, "6.1 调试方法", 2)
    add_body(doc, "调试采用编译检查、断点观察、异常堆栈、数据库数据核对和发布版本复测相结合的方法。每次修改先确认问题可复现，再缩小到具体页面或数据操作，修复后执行相关功能回归，避免只消除表面现象。")
    add_heading(doc, "6.2 主要问题与处理", 2)
    add_table(doc, ["问题", "原因分析", "处理方法", "验证结果"], [
        ("任务管理点击后卡死退出", "筛选控件初始化时提前触发 SelectionChanged，相关对象尚未就绪", "增加页面就绪标志，完成初始化后再允许刷新", "页面可稳定打开并筛选"),
        ("部分 XAML 资源加载失败", "资源后缀误写为 .xmal", "统一修正引用并重新编译", "资源正常加载"),
        ("SQLite 依赖安全告警", "底层原生包版本过旧", "升级 SQLitePCLRaw.lib.e_sqlite3 到安全版本", "构建无相关警告"),
        ("侧栏收起后文字残留", "导航项缺少明确图标与紧凑显示设计", "为每个模块配置图标并调整 Pane 行为", "收起后显示清晰图标"),
        ("应用只能在 Rider 中运行", "此前仅使用开发构建输出", "执行自包含 x64 发布并整理资源", "EXE 可独立启动"),
        ("品牌图标显示不一致", "窗口和打包资源未统一", "生成多尺寸资源并应用千里马图标", "窗口、启动页与发布包一致"),
    ], [3.4, 5.2, 5.4, 2.5])
    add_heading(doc, "6.3 测试结果", 2)
    add_table(doc, ["测试项", "操作", "预期结果", "实际结果"], [
        ("项目管理", "新增、编辑和删除项目", "数据正确更新", "通过"),
        ("任务闭环", "创建并推进至已完成", "状态按顺序变化", "通过"),
        ("筛选搜索", "组合关键词、项目和状态", "只显示符合条件记录", "通过"),
        ("逾期判断", "设置已过期且未完成任务", "显示逾期提示", "通过"),
        ("周报导出", "生成并导出 TXT", "文件内容完整", "通过"),
        ("数据备份", "执行一键备份", "生成可识别数据库副本", "通过"),
        ("独立运行", "双击发布目录 EXE", "正常进入主页面", "通过"),
    ], [3, 5, 6, 2.5])

    add_heading(doc, "7 系统总结", 1)
    add_heading(doc, "7.1 完成情况", 2)
    add_body(doc, "TeamFlowDesk 已完成项目结项所需的主要业务闭环。系统能够统一管理项目、任务、成员和设备，通过任务状态推进形成执行与验收流程；工作台提供整体视图；周报、导出和数据库备份支持成果沉淀；AI 协作记录体现开发中的判断与修正；自包含 EXE 解决了脱离开发环境运行的问题。")
    add_heading(doc, "7.2 系统特点", 2)
    add_list(doc, ["以任务流转为核心，而不是只完成静态增删改查。", "本地部署轻量，数据可备份，适合小型实训团队。", "导航支持展开与紧凑两种状态，兼顾信息量和可用空间。", "千里马品牌图标、开屏动画和统一视觉增强了产品完整性。", "AI 建议与人工验证分别记录，能够回溯开发决策。"])
    add_heading(doc, "7.3 客观不足", 2)
    add_body(doc, "受实训周期和桌面单机定位限制，系统当前以单机数据管理为主，不包含多设备实时同步和复杂权限系统；数据规模与并发场景也未达到企业级压力测试范围。这些限制不会影响当前面向小型团队的完整工作流，但明确了系统适用边界。")

    add_heading(doc, "8 课程设计总结", 1)
    add_heading(doc, "8.1 实践收获", 2)
    add_body(doc, "本次综合实践让我从单一功能编码转向完整产品交付。除了掌握 WinUI 3 页面、数据绑定、SQLite 持久化和 .NET 发布流程，我还进一步理解了需求边界、状态建模、异常处理、回归测试和文档表达的重要性。一个可用系统不仅要能运行，还要让用户理解当前状态、避免错误操作，并在故障发生时给出可处理的信息。")
    add_heading(doc, "8.2 AI 协作与批判性思维", 2)
    add_body(doc, "AI 在项目中主要承担方案生成、代码检查和文档整理，但其输出并不直接作为最终答案。对关键建议，我采用多角度验证：首先检查是否符合项目现有技术栈；其次评估改动范围和潜在回归；然后通过编译、运行和数据结果验证。")
    add_table(doc, ["AI 建议或问题", "人工质疑与验证", "最终处理"], [
        ("任务页崩溃可能来自数据库", "查看触发时机后发现尚未执行数据库查询，真正原因是控件初始化事件", "使用页面就绪标志阻止过早刷新"),
        ("通过隐藏文字解决侧栏收起问题", "隐藏文字不能保证模块可识别，且不符合紧凑导航习惯", "为每个导航项增加语义明确的图标"),
        ("直接复制数据库即可备份", "验证运行时文件占用、目标目录和重名问题", "关闭短连接后按时间戳复制，并返回结果提示"),
        ("修改框架版本可消除依赖警告", "框架升级可能扩大改动，检查后确认警告来自具体原生包", "只升级对应 SQLitePCLRaw 包并重新构建"),
    ], [4.6, 7.1, 5.8])
    add_body(doc, "这一过程说明，AI 更适合作为提供候选答案的协作者，而不是替代判断的工具。只有把建议放回具体代码、需求和运行结果中检查，才能形成可靠结论。")
    add_heading(doc, "8.3 个人总结", 2)
    add_body(doc, "项目最终从可运行原型发展为具有统一界面、完整工作流、数据保护和独立发布能力的桌面应用。最有价值的经验是先明确真实问题，再选择最小且可靠的修改，并在修复后进行相关功能回归。通过本次实践，我具备了从需求分析、数据库设计、界面实现到调试发布和报告整理的较完整开发经历。")

    add_heading(doc, "9 参考文献", 1)
    refs = [
        "[1] Microsoft. Windows App SDK Documentation. Microsoft Learn.",
        "[2] Microsoft. WinUI 3 Documentation. Microsoft Learn.",
        "[3] Microsoft. .NET Documentation. Microsoft Learn.",
        "[4] SQLite Consortium. SQLite Documentation.",
        "[5] 萨师煊, 王珊. 数据库系统概论. 高等教育出版社.",
        "[6] 张海藩. 软件工程导论. 清华大学出版社.",
        "[7] Ian Sommerville. Software Engineering. Pearson.",
    ]
    for ref in refs:
        add_body(doc, ref, first=False)

    # Prevent heading orphans and set metadata.
    doc.core_properties.title = "TeamFlowDesk 团队运行管理工作流系统 软件综合实践报告"
    doc.core_properties.author = "杨杨文琦"
    doc.core_properties.subject = "软件综合实践结项报告"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
