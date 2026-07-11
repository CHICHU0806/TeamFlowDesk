from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "outputs" / "软件综合实践报告-保留模板版.docx"
DST = ROOT / "outputs" / "软件综合实践报告-人工智能1班-20240282-杨杨文琦-按原模板填写版.docx"
FIG = ROOT / "outputs" / "report_figures"
ICON = ROOT / "Assets" / "TeamFlowDeskIcon.png"


def style_run(run, size=10.5, bold=False, name="宋体", color="000000"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def insert_p_after(anchor, text="", subheading=False, first=True, center=False):
    p = OxmlElement("w:p")
    anchor.addnext(p)
    para = anchor.getparent().getparent().part.document.paragraphs[-1] if False else None
    # Construct through a temporary document paragraph, then move its XML node.
    doc = anchor.getroottree().getroot()
    return p


def make_para(doc, text, subheading=False, first=True, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(5 if subheading else 0)
    p.paragraph_format.space_after = Pt(5)
    if first and not subheading and not center:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    style_run(r, size=12 if subheading else 10.5, bold=subheading, name="黑体" if subheading else "宋体", color="1F4E79" if subheading else "000000")
    return p


def after(doc, anchor, text, subheading=False, first=True, center=False):
    p = make_para(doc, text, subheading, first, center)
    anchor.addnext(p._p)
    return p._p


def add_table_after(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "D9EAF7"); c._tc.get_or_add_tcPr().append(shd)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(c.paragraphs[0].add_run(h), 9.5, True, "黑体", "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Cm(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_run(cells[i].paragraphs[0].add_run(str(value)), 9)
    anchor.addnext(table._tbl)
    return table._tbl


def add_picture_after(doc, anchor, path, width, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    anchor.addnext(p._p)
    cap = make_para(doc, caption, first=False, center=True)
    for run in cap.runs:
        style_run(run, 9, False, "宋体", "555555")
    p._p.addnext(cap._p)
    return cap._p


def find_para(doc, exact):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p._p
    raise ValueError(exact)


def build():
    doc = Document(SRC)

    # Only fill blank/intended fields; existing template text and dates remain untouched.
    style_run(doc.tables[0].cell(0, 1).paragraphs[0].add_run("杨杨文琦"), 12)
    style_run(doc.tables[0].cell(1, 1).paragraphs[0].add_run("20240282"), 12)
    style_run(doc.paragraphs[10].add_run("TeamFlowDesk 团队运行管理工作流系统"), 12, True, "黑体")

    # 1 绪论
    a = find_para(doc, "主要对系统的背景进行简要的说明。")
    a = after(doc, a, "1.1 项目背景", True)
    a = after(doc, a, "在软件综合实践中，小型团队往往同时使用聊天记录、表格和纸面记录管理项目，容易造成任务信息分散、负责人不明确、截止时间遗漏以及阶段成果难以汇总。为解决这些问题，本项目开发 TeamFlowDesk 团队运行管理工作流系统，将项目、任务、成员、设备、周报和 AI 协作记录集中到一个 Windows 桌面应用中。")
    a = after(doc, a, "1.2 项目目标", True)
    a = after(doc, a, "系统以任务闭环为核心，支持任务从待处理、进行中、待验收到已完成的全过程管理，并通过项目关联、负责人、优先级、截止日期、搜索筛选和逾期提醒提升执行效率。同时提供周报导出、数据库备份和自包含 EXE 发布，使应用具备完整的使用和交付条件。")
    a = after(doc, a, "1.3 技术路线", True)
    a = after(doc, a, "项目采用 C#、.NET 10、WinUI 3、Windows App SDK 和 SQLite 开发。界面层负责导航、表单和数据展示，业务逻辑负责状态校验和统计，数据访问层负责参数化 SQL 与本地持久化。开发过程按照需求分析、原型实现、功能补齐、异常修复、界面优化和发布验证逐步推进。")

    # 2 需求分析
    a = find_para(doc, "需求描述，系统功能模块图，主要功能描述，用例图。")
    a = after(doc, a, "2.1 需求描述", True)
    a = after(doc, a, "系统面向 3 至 10 人的小型实训团队。项目负责人需要掌握整体进度、任务逾期和成员分工；普通成员需要查看自己的待办及任务状态；团队还需要管理共用设备并形成阶段周报。系统应操作直观、响应及时、数据可保存和备份，并可脱离开发工具独立运行。")
    a = add_table_after(doc, a, ["模块", "主要功能", "完成状态"], [
        ("工作台", "展示项目、任务、成员、逾期数量和近期动态", "已完成"),
        ("项目管理", "项目新增、编辑、删除、状态和进度维护", "已完成"),
        ("任务管理", "搜索筛选、项目关联、负责人、截止日期和状态流转", "已完成"),
        ("成员管理", "成员档案、角色和任务负载查看", "已完成"),
        ("设备管理", "设备登记、状态和借用信息维护", "已完成"),
        ("AI 协作", "记录 AI 建议、人工验证和修正结果", "已完成"),
        ("周报与设置", "周报导出、数据库备份和关于信息", "已完成"),
    ], [3.2, 10.2, 2.5])
    a = after(doc, a, "2.2 功能模块图", True)
    a = add_picture_after(doc, a, FIG / "modules.png", 15.5, "图2-1 系统功能模块图")
    a = after(doc, a, "2.3 主要用例", True)
    a = after(doc, a, "创建任务用例：用户进入任务管理页，填写任务标题、所属项目、负责人、优先级和截止时间。系统校验输入后写入数据库，新任务默认为“待处理”。负责人开始执行后将其推进到“进行中”，提交成果后进入“待验收”，验收通过后转为“已完成”。若验收未通过，可退回上一状态修正。")
    a = after(doc, a, "生成周报用例：用户选择统计周期，系统汇总周期内的项目、任务和成员数据，形成周报内容并导出 TXT 文件。备份用例：用户在设置页执行备份，系统将当前 SQLite 数据库复制到文档目录的 TeamFlowDeskBackups 文件夹，并返回保存结果。")

    # 3 数据库设计: keep all four instruction paragraphs and insert after the last.
    a = find_para(doc, "数据库物理设计（用SQL语言实现），生成创建数据库的脚本，包括数据库结构定义 ，建立索引、视图语句，存储过程（如果使用）的结构和定义，主要的查询语句等。")
    a = after(doc, a, "3.1 数据流与概念结构", True)
    a = after(doc, a, "用户在页面输入业务数据，界面完成格式校验后调用业务逻辑，再由数据访问层以参数化 SQL 写入 SQLite。工作台和周报中心从多个业务表读取数据并进行聚合。核心实体为项目、任务、成员、设备、AI 协作记录和周报；一个项目包含多个任务，一个成员可负责多个任务，设备借用信息可关联成员。")
    a = add_table_after(doc, a, ["数据表", "关键字段", "用途"], [
        ("Projects", "Id、Name、Status、Progress、StartDate、EndDate", "保存项目基本信息和进度"),
        ("Tasks", "ProjectId、Title、AssigneeId、Priority、Status、DueDate", "保存任务归属、负责人和状态"),
        ("Members", "Name、Role、Email、Phone、Status", "保存成员档案"),
        ("Equipment", "Name、Code、Status、BorrowerId、BorrowDate", "保存设备及借用信息"),
        ("AiRecords", "Question、AiSuggestion、HumanDecision、Correction", "保存 AI 协作和人工修正"),
        ("WeeklyReports", "WeekStart、WeekEnd、Summary、CreatedAt", "保存周报内容"),
    ], [3, 8.2, 4.8])
    a = after(doc, a, "3.2 逻辑与物理设计", True)
    a = after(doc, a, "各表使用自增整数主键。Tasks.ProjectId 关联 Projects.Id，Tasks.AssigneeId 关联 Members.Id。日期统一保存为可排序文本，任务状态限定为待处理、进行中、待验收和已完成，优先级限定为低、中、高。系统删除项目时先检查关联任务，防止产生失去归属的数据。所有写入语句均使用参数传值。")
    a = after(doc, a, "3.3 主要 SQL 实现", True)
    sql = "CREATE TABLE IF NOT EXISTS Tasks (Id INTEGER PRIMARY KEY AUTOINCREMENT, ProjectId INTEGER, Title TEXT NOT NULL, Description TEXT, AssigneeId INTEGER, Priority TEXT NOT NULL, Status TEXT NOT NULL, DueDate TEXT, CreatedAt TEXT NOT NULL, UpdatedAt TEXT NOT NULL, FOREIGN KEY(ProjectId) REFERENCES Projects(Id), FOREIGN KEY(AssigneeId) REFERENCES Members(Id));"
    a = after(doc, a, sql, first=False)
    a = after(doc, a, "常用查询包括按项目、状态和关键词组合筛选任务，统计未完成与逾期任务数量，以及按负责人汇总任务负载。SQLite 适合本项目的单机小团队场景，部署简单且数据库文件便于备份。")

    # 4 系统设计
    a = find_para(doc, "类图、时序图、活动图等，架构。")
    a = after(doc, a, "4.1 系统架构", True)
    a = add_picture_after(doc, a, FIG / "architecture.png", 15.5, "图4-1 TeamFlowDesk 分层架构图")
    a = after(doc, a, "表示层由 MainWindow 和各业务页面组成；业务层处理状态规则、统计、导出和备份；数据访问层统一管理 SQLite 连接及查询；数据层保存业务实体。分层设计降低了界面、业务规则和数据库之间的耦合。")
    a = after(doc, a, "4.2 任务活动图", True)
    a = add_picture_after(doc, a, FIG / "workflow.png", 15.7, "图4-2 任务状态活动图")
    a = after(doc, a, "4.3 类与时序说明", True)
    a = after(doc, a, "MainWindow 负责启动动画、标题栏、紧凑导航和页面切换；DashboardPage、TasksPage 等页面负责展示数据和采集操作；Project、TaskItem、Member 等模型描述业务实体；DatabaseService 负责初始化和持久化。以创建任务为例，TasksPage 校验输入后调用数据服务执行 INSERT，成功后重新加载列表与统计，失败则捕获异常并显示原因。")

    # 5 编码实现
    a = find_para(doc, "开发环境，程序的运行环境，主要运行界面的截图。")
    a = after(doc, a, "5.1 开发与运行环境", True)
    a = add_table_after(doc, a, ["项目", "配置"], [("开发语言", "C#"), ("运行框架", ".NET 10"), ("桌面框架", "WinUI 3 / Windows App SDK"), ("数据库", "SQLite"), ("开发工具", "JetBrains Rider"), ("运行平台", "Windows 10/11 x64"), ("发布形式", "Self-contained 独立 EXE")], [4.2, 12])
    a = after(doc, a, "5.2 主要功能实现", True)
    a = after(doc, a, "主窗口使用 NavigationView 组织功能。侧栏展开时显示图标与完整文字，收起后仅显示语义明确的模块图标。左上角、标题栏和发布资源统一使用千里马头图标。应用启动时图标逐渐显现、短暂停留后淡出，再进入工作台，动画总时长约 1.25 秒。")
    a = add_picture_after(doc, a, ICON, 4.2, "图5-1 TeamFlowDesk 千里马专用图标")
    a = after(doc, a, "任务页提供关键词搜索、项目和状态筛选、任务详情、截止日期预警以及状态推进。工作台聚合项目数、进行中任务数、逾期任务数和成员数。成员页展示成员档案与负载，设备页维护设备编码和借用状态，AI 协作页记录问题、建议、验证与纠正，周报中心可生成并导出 TXT，设置页可备份数据库。")
    a = after(doc, a, "5.3 独立发布", True)
    a = after(doc, a, "项目以 Windows x64 自包含方式发布，用户可以直接双击 TeamFlowDesk.exe 启动，不依赖 Rider，也不要求目标电脑预装对应 .NET 运行时。发布后对启动动画、页面导航、任务状态流转、数据库读写、导出和备份进行了复测。")

    # 6 调试
    a = find_para(doc, "关键代码分析，系统在调试过程中出现问题都要加以说明。")
    a = after(doc, a, "6.1 调试方法", True)
    a = after(doc, a, "调试采用编译检查、断点观察、异常堆栈、数据库核对和发布版本复测相结合的方法。每个问题先确认可复现，再定位到事件、页面或数据操作，修复后执行相关功能回归。")
    a = add_table_after(doc, a, ["问题", "原因", "处理与结果"], [
        ("点击任务管理后卡死退出", "筛选控件初始化时过早触发 SelectionChanged", "增加页面就绪标志，初始化完成后再刷新，问题解决"),
        ("部分 XAML 资源加载失败", "资源后缀误写为 .xmal", "统一修正为 .xaml，资源恢复正常"),
        ("SQLite 依赖安全告警", "底层原生包版本过旧", "升级 SQLitePCLRaw.lib.e_sqlite3，构建无相关警告"),
        ("侧栏收起后显示半截文字", "导航项缺少紧凑状态图标", "为所有模块配置图标，收起后正常缩略显示"),
        ("只能从 Rider 运行", "尚未进行自包含发布", "生成独立 EXE，双击可正常启动"),
    ], [4.1, 5.5, 7])
    a = after(doc, a, "6.2 关键代码分析", True)
    a = after(doc, a, "任务页使用 _isPageReady 标志区分控件初始化事件与真实用户操作。页面加载完成前，筛选变化事件直接返回；项目、成员和筛选数据准备完成后再统一刷新任务列表。该修复避免了空对象访问和重复查询，同时不影响用户后续筛选。数据写入采用参数化 SQL，异常由页面捕获并反馈，不让数据库错误直接导致程序退出。")

    # 7 系统总结
    a = find_para(doc, "对本系统作一个全面的评价：包括有何特点、存在的问题、改进意见等。")
    a = after(doc, a, "7.1 系统评价", True)
    a = after(doc, a, "TeamFlowDesk 已形成从项目建立、任务分配、执行、验收到统计归档的完整闭环。系统不仅实现常规增删改查，还加入任务状态规则、搜索筛选、逾期提醒、成员负载、周报导出、数据库备份和独立发布，能够满足小型实训团队的日常管理需求。")
    a = after(doc, a, "7.2 系统特点", True)
    a = after(doc, a, "系统以任务流转作为核心业务，而不是将任务状态作为静态文字；本地 SQLite 部署轻量且便于备份；展开与紧凑导航兼顾信息完整性和可用空间；千里马品牌图标、开屏动画和统一视觉提升了产品完整度；AI 协作记录将建议与人工验证分开保存，便于复盘。")
    a = after(doc, a, "7.3 存在的问题及改进意见", True)
    a = after(doc, a, "受实训周期和桌面单机定位限制，系统当前主要服务于本机用户，不具备多设备实时同步和复杂权限控制；数据规模与并发测试也未达到企业级范围。若应用场景扩大，可在保持现有桌面交互的基础上增加服务端同步、账号权限和自动化测试。上述内容属于适用边界说明，不影响当前项目的完整运行。")

    # 8 课程设计总结
    a = find_para(doc, "通过做课程设计，你有何感想，学到了什么？")
    a = after(doc, a, "通过本次综合实践，我从实现单一页面功能转向完成一个可交付的软件产品。除 WinUI 3 页面、数据绑定、SQLite 持久化和 .NET 发布外，我还学习了需求边界、状态建模、异常处理、回归测试和技术文档表达。一个可用系统不仅要能够运行，还要让用户理解当前状态、避免误操作，并在发生异常时提供清晰反馈。")
    a = after(doc, a, "AI 在项目中用于生成候选方案、检查代码和整理文档，但我没有直接采用其全部结论。针对任务页崩溃，我先质疑“数据库故障”的判断，通过事件触发顺序确认真正原因是控件尚未初始化；针对依赖告警，我没有盲目升级整个框架，而是定位到具体 SQLite 原生包；针对侧栏问题，我否定了单纯隐藏文字的做法，改用模块图标保证可识别性。每项关键修改均通过编译、运行和数据结果进行验证。")
    a = after(doc, a, "本次实践使我认识到，AI 更适合作为提供候选答案的协作者，而不是替代判断的工具。可靠的软件开发需要将建议放回具体需求、代码结构和运行结果中核验。项目最终从可运行原型发展为具有统一界面、完整工作流、数据保护和独立发布能力的桌面应用，我也获得了从需求分析到结项交付的完整实践经验。")

    # 9 参考文献: insert without changing heading.
    a = find_para(doc, "参考文献")
    refs = [
        "[1] Microsoft. Windows App SDK Documentation. Microsoft Learn.",
        "[2] Microsoft. WinUI 3 Documentation. Microsoft Learn.",
        "[3] Microsoft. .NET Documentation. Microsoft Learn.",
        "[4] SQLite Consortium. SQLite Documentation.",
        "[5] 王珊, 萨师煊. 数据库系统概论. 高等教育出版社.",
        "[6] 张海藩. 软件工程导论. 清华大学出版社.",
    ]
    for ref in refs:
        a = after(doc, a, ref, first=False)

    doc.core_properties.title = "TeamFlowDesk 软件综合实践报告（按原模板填写）"
    doc.core_properties.author = "杨杨文琦"
    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    build()
